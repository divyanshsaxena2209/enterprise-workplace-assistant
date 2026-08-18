from docx import Document
document = Document()
document.add_heading('Test Document', 0)
p = document.add_paragraph('This is a test document.')
document.save('test.docx')
